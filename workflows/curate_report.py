import os
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
import yfinance as yf
import pandas as pd
from datetime import datetime, timedelta
from rich.console import Console
from rich.panel import Panel
from rich.markdown import Markdown as RichMarkdown
from state import FinanceAgentState
from langchain_core.messages import AIMessage, HumanMessage 

# Try to import docx with fallback
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    DOCX_AVAILABLE = True
    print("✅ python-docx library loaded successfully")
except ImportError as e:
    print(f"⚠ Warning: python-docx not available: {e}")
    DOCX_AVAILABLE = False

# Initialize Rich console for pretty printing
console = Console()

def create_price_chart(company_name, symbol=None, days=90):
    """Create a professional price chart with gradient styling"""
    try:
        if not symbol:
            symbol_map = {
                'salesforce': 'CRM',
                'apple': 'AAPL',
                'microsoft': 'MSFT',
                'google': 'GOOGL',
                'tesla': 'TSLA',
                'amazon': 'AMZN',
                'meta': 'META',
                'nvidia': 'NVDA',
                'tata motors':'TATMOT', 
            }
            symbol = symbol_map.get(company_name.lower().split()[0], company_name.upper()[:4])
        
        end_date = datetime.now()
        start_date = end_date - timedelta(days=days)
        
        console.print(f"📊 Downloading stock data for {symbol}...")
        
        stock_data = yf.download(
            symbol, 
            start=start_date, 
            end=end_date, 
            progress=False,
            auto_adjust=False,  
            threads=True
        )
        
        if stock_data.empty:
            console.print(Panel(f"⚠️ No stock data found for {symbol}", style="bold yellow"))
            return None, None, None
        
        if isinstance(stock_data.index, pd.MultiIndex):
            stock_data = stock_data.droplevel(1, axis=1)
        
        stock_data = stock_data.reset_index()
        if 'Date' not in stock_data.columns and stock_data.index.name == 'Date':
            stock_data = stock_data.reset_index()
        
        required_columns = ['Close']
        if not all(col in stock_data.columns for col in required_columns):
            console.print(Panel(f"⚠️ Missing required data columns for {symbol}", style="bold yellow"))
            return None, None, None
        
        # Enhanced chart styling
        plt.style.use('seaborn-v0_8-darkgrid')
        fig, ax = plt.subplots(figsize=(14, 7), facecolor='#f8f9fa')
        
        if 'Date' in stock_data.columns:
            dates = pd.to_datetime(stock_data['Date'])
        else:
            dates = pd.to_datetime(stock_data.index)
        
        prices = stock_data['Close']
        
        # Gradient fill effect
        ax.plot(dates, prices, color="#1f77b4", linewidth=3, alpha=0.9, label=symbol)
        ax.fill_between(dates, prices, alpha=0.2, color='#4CAF50')
        
        # Enhanced grid and background
        ax.set_facecolor("#ffffff")
        ax.grid(True, alpha=0.2, linestyle='--', linewidth=0.8, color='#cccccc')
        
        ax.xaxis.set_major_formatter(mdates.DateFormatter('%b %d'))
        ax.xaxis.set_major_locator(mdates.WeekdayLocator(interval=2))
        plt.xticks(rotation=45, fontsize=10)
        
        ax.set_title(f'{company_name} Stock Price Performance - Last {days} Days', 
                    fontsize=18, fontweight='bold', color='#1a1a1a', pad=20)
        ax.set_xlabel('Date', fontsize=13, color='#333333', fontweight='600')
        ax.set_ylabel('Price ($)', fontsize=13, color='#333333', fontweight='600')
        
        # Enhanced price annotation
        current_price = prices.iloc[-1]
        latest_date = dates.iloc[-1]
        ax.annotate(f'${current_price:.2f}', 
                   xy=(latest_date, current_price),
                   xytext=(15, 15), textcoords='offset points',
                   bbox=dict(boxstyle='round,pad=0.5', facecolor='#1f77b4', 
                            edgecolor='#0d47a1', linewidth=2, alpha=0.9),
                   color='white', fontweight='bold', fontsize=11,
                   arrowprops=dict(arrowstyle='->', color='#1f77b4', lw=2))
        
        # Add legend
        ax.legend(loc='upper left', fontsize=10, framealpha=0.9)
        
        plt.tight_layout()
        
        chart_dir = './reports/charts'
        os.makedirs(chart_dir, exist_ok=True)
        chart_path = os.path.join(chart_dir, f'{symbol}_chart_{datetime.now().strftime("%Y%m%d_%H%M%S")}.png')
        
        plt.savefig(chart_path, dpi=300, bbox_inches='tight', facecolor='#f8f9fa')
        plt.close()
        
        console.print(Panel(f"✅ Chart saved: {chart_path}", style="bold green"))
        return chart_path, current_price, symbol
        
    except Exception as e:
        console.print(Panel(f"⚠️ Chart creation failed: {str(e)}", style="bold yellow"))
        import traceback
        console.print(f"Chart error traceback: {traceback.format_exc()}")
        return None, None, None

def add_cell_background_color(cell, color_hex):
    """Helper function to add background color to table cells"""
    try:
        shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
        cell._tc.get_or_add_tcPr().append(shading_elm)
        return True
    except Exception as e:
        console.print(f"Warning: Could not set cell background color: {e}")
        return False

def add_gradient_header(doc):
    """Add a beautiful gradient header section"""
    try:
        header_section = doc.sections[0]
        header = header_section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True
    except:
        return False

def truncate_text(text, max_length=2000):
    """Truncate text to a maximum length with ellipsis"""
    if not text or len(str(text)) <= max_length:
        return str(text) if text else ""
    return str(text)[:max_length] + "...\n\n[Content truncated for brevity. See full filing for complete details.]"

def create_beautiful_risk_section(doc, risk_summary):
    """Create a visually striking risk section with warning styling"""
    
    # Risk section header with icon and color
    risk_header = doc.add_paragraph()
    risk_header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Warning icon and title
    icon_run = risk_header.add_run('⚠️ ')
    icon_run.font.size = Pt(20)
    
    title_run = risk_header.add_run('KEY BUSINESS RISKS')
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(220, 53, 69)  # Bootstrap danger red
    
    subtitle_para = doc.add_paragraph()
    subtitle_run = subtitle_para.add_run('From SEC 10-K Filing - Item 1A Risk Factors')
    subtitle_run.font.name = 'Arial'
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.italic = True
    subtitle_run.font.color.rgb = RGBColor(108, 117, 125)
    
    # Create a bordered box effect with table
    risk_table = doc.add_table(rows=1, cols=1)
    risk_table.style = 'Table Grid'
    risk_cell = risk_table.rows[0].cells[0]
    
    # Add colored border
    add_cell_background_color(risk_cell, 'FFF3CD')  # Light warning yellow
    
    # Add risk content
    if risk_summary and str(risk_summary).strip() and str(risk_summary).lower() not in ["none", "n/a"]:
        truncated_risk = truncate_text(risk_summary, max_length=3000)
        risk_para = risk_cell.paragraphs[0]
        risk_text = risk_para.add_run(truncated_risk)
        risk_text.font.name = 'Arial'
        risk_text.font.size = Pt(10)
        risk_text.font.color.rgb = RGBColor(51, 51, 51)
    else:
        risk_para = risk_cell.paragraphs[0]
        risk_text = risk_para.add_run(
            "⚠️ Risk factor analysis data is currently unavailable. "
            "Please refer to the company's latest 10-K filing Item 1A (Risk Factors) "
            "for detailed information about business risks."
        )
        risk_text.font.name = 'Arial'
        risk_text.font.size = Pt(10)
        risk_text.font.italic = True
        risk_text.font.color.rgb = RGBColor(108, 117, 125)
    
    doc.add_paragraph()  # Spacing

def create_beautiful_metrics_table(doc, key_metrics):
    """Create a stunning metrics table with alternating row colors"""
    
    metrics_header = doc.add_paragraph()
    icon_run = metrics_header.add_run('💰 ')
    icon_run.font.size = Pt(18)
    
    title_run = metrics_header.add_run('KEY FINANCIAL METRICS')
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(40, 167, 69)  # Bootstrap success green
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light List Accent 1'
    
    # Enhanced header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Metric'
    header_cells[1].text = 'Value'
    
    for cell in header_cells:
        add_cell_background_color(cell, '007BFF')  # Bootstrap primary blue
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Arial'
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Add formatted metrics with alternating colors
    metric_labels = {
        'current_price': '💵 Current Price',
        'market_cap': '🏢 Market Cap',
        'pe_ratio': '📊 P/E Ratio',
        'eps': '💹 EPS',
        'week_52_high': '⬆️ 52-Week High',
        'week_52_low': '⬇️ 52-Week Low'
    }
    
    row_colors = ['F8F9FA', 'FFFFFF']  # Alternating light gray and white
    added_rows = 0
    
    for key, value in key_metrics.items():
        if value is not None and value != 'N/A' and str(value).strip():
            try:
                row = table.add_row().cells
                row[0].text = metric_labels.get(key, key.replace('_', ' ').title())
                
                # Format value based on type
                if key == 'market_cap' and isinstance(value, (int, float)) and value > 0:
                    if value >= 1e12:
                        formatted_value = f"${value/1e12:.2f}T"
                    elif value >= 1e9:
                        formatted_value = f"${value/1e9:.2f}B"
                    else:
                        formatted_value = f"${value/1e6:.2f}M"
                elif key in ['current_price', 'week_52_high', 'week_52_low'] and isinstance(value, (int, float)):
                    formatted_value = f"${value:.2f}"
                else:
                    formatted_value = str(value)
                
                row[1].text = formatted_value
                
                # Apply alternating colors
                color = row_colors[added_rows % 2]
                for cell in row:
                    add_cell_background_color(cell, color)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Arial'
                            run.font.size = Pt(10)
                            # Make value bold
                            if cell == row[1]:
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(0, 123, 255)
                
                added_rows += 1
            except Exception as e:
                console.print(f"⚠️ Could not add metric {key}: {e}")
                continue
    
    if added_rows > 0:
        doc.add_paragraph()
        console.print(f"✅ Added {added_rows} metrics to table")

def create_enhanced_docx_report(company_name, risk_summary, stock_performance_text, news_articles, key_metrics):
    """Create a stunning professional DOCX report with modern design"""
    if not DOCX_AVAILABLE:
        console.print(Panel("❌ Cannot create DOCX: python-docx not available", style="bold red"))
        return None
    
    try:
        console.print("📄 Initializing beautiful document...")
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # ========== STUNNING HEADER SECTION ==========
        # Title with gradient effect (simulated with colors)
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        title_run = title_para.add_run('📊 FINANCIAL ANALYSIS REPORT')
        title_run.font.name = 'Arial'
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(13, 71, 161)  # Deep blue
        
        # Company name with accent color
        company_para = doc.add_paragraph()
        company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        company_run = company_para.add_run(str(company_name))
        company_run.font.name = 'Arial'
        company_run.font.size = Pt(24)
        company_run.font.bold = True
        company_run.font.color.rgb = RGBColor(220, 53, 69)  # Accent red
        
        # Elegant date stamp
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(f'Generated: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}')
        date_run.font.name = 'Arial'
        date_run.font.size = Pt(10)
        date_run.font.italic = True
        date_run.font.color.rgb = RGBColor(108, 117, 125)
        
        # Decorative separator
        separator_para = doc.add_paragraph()
        separator_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        separator_run = separator_para.add_run('━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━')
        separator_run.font.color.rgb = RGBColor(13, 71, 161)
        
        doc.add_paragraph()
        
        # ========== CHART SECTION ==========
        console.print("📈 Creating and inserting chart...")
        chart_path, current_price, symbol = create_price_chart(company_name)
        if chart_path and os.path.exists(chart_path):
            try:
                chart_para = doc.add_paragraph()
                chart_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                chart_run = chart_para.add_run()
                chart_run.add_picture(chart_path, width=Inches(6.5))
                doc.add_paragraph()
                console.print("✅ Chart successfully added to document")
            except Exception as e:
                console.print(f"⚠️ Could not add chart to document: {e}")
        
        # ========== EXECUTIVE SUMMARY ==========
        exec_header = doc.add_paragraph()
        exec_icon = exec_header.add_run('📋 ')
        exec_icon.font.size = Pt(18)
        
        exec_title = exec_header.add_run('EXECUTIVE SUMMARY')
        exec_title.font.name = 'Arial'
        exec_title.font.size = Pt(18)
        exec_title.font.bold = True
        exec_title.font.color.rgb = RGBColor(52, 58, 64)
        
        exec_body = doc.add_paragraph()
        exec_text = exec_body.add_run(
            'This comprehensive financial analysis synthesizes insights from SEC filings, '
            'real-time market data, and current news sentiment to provide actionable '
            'investment intelligence. The report highlights key financial metrics, material '
            'business risks, and recent market developments.'
        )
        exec_text.font.name = 'Arial'
        exec_text.font.size = Pt(11)
        exec_text.font.color.rgb = RGBColor(51, 51, 51)
        
        doc.add_paragraph()
        
        # ========== KEY METRICS TABLE ==========
        if key_metrics and isinstance(key_metrics, dict) and any(key_metrics.values()):
            console.print("📊 Adding beautiful metrics table...")
            create_beautiful_metrics_table(doc, key_metrics)
        
        # ========== RISK ANALYSIS SECTION ==========
        console.print("⚠️ Adding enhanced risk section...")
        create_beautiful_risk_section(doc, risk_summary)
        
        # ========== MARKET PERFORMANCE ==========
        perf_header = doc.add_paragraph()
        perf_icon = perf_header.add_run('📈 ')
        perf_icon.font.size = Pt(18)
        
        perf_title = perf_header.add_run('MARKET PERFORMANCE SUMMARY')
        perf_title.font.name = 'Arial'
        perf_title.font.size = Pt(18)
        perf_title.font.bold = True
        perf_title.font.color.rgb = RGBColor(40, 167, 69)
        
        perf_body = doc.add_paragraph()
        clean_performance_text = str(stock_performance_text).replace('**', '') if stock_performance_text else "Performance data unavailable."
        perf_text = perf_body.add_run(clean_performance_text)
        perf_text.font.name = 'Arial'
        perf_text.font.size = Pt(11)
        perf_text.font.color.rgb = RGBColor(51, 51, 51)
        
        doc.add_paragraph()
        
        # ========== NEWS SECTION ==========
        news_header = doc.add_paragraph()
        news_icon = news_header.add_run('📰 ')
        news_icon.font.size = Pt(18)
        
        news_title = news_header.add_run('RECENT NEWS & MARKET CATALYSTS')
        news_title.font.name = 'Arial'
        news_title.font.size = Pt(18)
        news_title.font.bold = True
        news_title.font.color.rgb = RGBColor(253, 126, 20)  # Orange
        
        if news_articles and isinstance(news_articles, list) and len(news_articles) > 0:
            news_intro = doc.add_paragraph()
            intro_text = news_intro.add_run('Recent market developments and news coverage:')
            intro_text.font.name = 'Arial'
            intro_text.font.size = Pt(11)
            intro_text.font.color.rgb = RGBColor(51, 51, 51)
            
            for i, article in enumerate(news_articles[:5], 1):
                if isinstance(article, dict) and article.get('title'):
                    news_item = doc.add_paragraph()
                    news_item.paragraph_format.left_indent = Inches(0.25)
                    
                    bullet_run = news_item.add_run(f'{i}. ')
                    bullet_run.font.name = 'Arial'
                    bullet_run.font.size = Pt(10)
                    bullet_run.font.bold = True
                    bullet_run.font.color.rgb = RGBColor(253, 126, 20)
                    
                    title = str(article.get('title', 'N/A'))
                    item_text = news_item.add_run(title)
                    item_text.font.name = 'Arial'
                    item_text.font.size = Pt(10)
                    item_text.font.color.rgb = RGBColor(51, 51, 51)
        else:
            news_body = doc.add_paragraph()
            news_text = news_body.add_run('Recent market developments and analyst coverage continue to evolve. Monitor financial news sources for the latest updates.')
            news_text.font.name = 'Arial'
            news_text.font.size = Pt(11)
            news_text.font.color.rgb = RGBColor(108, 117, 125)
        
        # ========== DISCLAIMER SECTION ==========
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Decorative separator before disclaimer
        separator_para2 = doc.add_paragraph()
        separator_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        separator_run2 = separator_para2.add_run('━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━')
        separator_run2.font.color.rgb = RGBColor(220, 53, 69)
        
        disclaimer_header = doc.add_paragraph()
        disclaimer_icon = disclaimer_header.add_run('⚖️ ')
        disclaimer_icon.font.size = Pt(16)
        
        disclaimer_title = disclaimer_header.add_run('IMPORTANT DISCLAIMER')
        disclaimer_title.font.name = 'Arial'
        disclaimer_title.font.size = Pt(14)
        disclaimer_title.font.bold = True
        disclaimer_title.font.color.rgb = RGBColor(220, 53, 69)
        
        # Disclaimer in bordered box
        disclaimer_table = doc.add_table(rows=1, cols=1)
        disclaimer_table.style = 'Table Grid'
        disclaimer_cell = disclaimer_table.rows[0].cells[0]
        add_cell_background_color(disclaimer_cell, 'F8D7DA')  # Light red background
        
        disclaimer_para = disclaimer_cell.paragraphs[0]
        disclaimer_text = disclaimer_para.add_run(
            'This report is generated for informational and educational purposes only. It should not be '
            'considered as personalized investment advice, a recommendation to buy or sell securities, '
            'or a guarantee of future performance. All investments carry risk of loss. Past performance '
            'does not guarantee future results. Please consult with a qualified financial advisor '
            'before making any investment decisions. The information contained in this report is based '
            'on publicly available data and may contain errors or omissions.'
        )
        disclaimer_text.font.name = 'Arial'
        disclaimer_text.font.size = Pt(9)
        disclaimer_text.font.color.rgb = RGBColor(51, 51, 51)
        
        console.print("✅ Beautiful document structure created successfully")
        return doc
        
    except Exception as e:
        console.print(Panel(f"❌ Error creating document: {str(e)}", style="bold red"))
        import traceback
        console.print(f"Document creation traceback: {traceback.format_exc()}")
        return None

def curate_report_node(state: FinanceAgentState) -> dict:
    """
    Synthesizes information from multiple sources to generate a comprehensive
    beautiful financial analyst report.
    """
    
    console.print(Panel("📄 NODE: Generating Beautiful Professional Report", style="bold green"))

    # --- 1. RETRIEVE DATA FROM STATE ---
    company_name = state.get("company_name", "Unknown Company")
    ticker = state.get("ticker", "")
    risk_summary = state.get("tool_result")
    filing_info = state.get("filing_info", {})
    
    # Debug logging
    console.print(f"[DEBUG] tool_result type: {type(risk_summary)}")
    console.print(f"[DEBUG] tool_result length: {len(str(risk_summary)) if risk_summary else 0}")
    console.print(f"[DEBUG] tool_result preview: {str(risk_summary)[:300] if risk_summary else 'None'}...")
    
    retrieved_section = filing_info.get("section", "Unknown")
    console.print(f"[DEBUG] Retrieved section: {retrieved_section}")
    
    # Validate risk summary
    if not risk_summary or str(risk_summary).strip() == "" or str(risk_summary).lower() in ["none", "n/a"]:
        console.print("⚠️ WARNING: Risk summary is empty or invalid!")
        if filing_info:
            filing_url = filing_info.get("filing_url", "N/A")
            risk_summary = (
                f"⚠️ Risk factor data could not be retrieved from {company_name}'s SEC filings.\n\n"
                f"Filing Information:\n"
                f"- Ticker: {ticker}\n"
                f"- Section Requested: {retrieved_section}\n"
                f"- Filing URL: {filing_url}\n\n"
                f"Please check the latest 10-K filing Item 1A for detailed risk information."
            )
        else:
            risk_summary = f"Risk factor data could not be retrieved from {company_name}'s SEC filings."
    else:
        console.print(f"✅ Risk summary retrieved: {len(str(risk_summary))} characters")
    
    key_metrics = state.get("structured_data", {})
    if not isinstance(key_metrics, dict):
        key_metrics = {}
    
    news_articles = state.get("news_results", [])
    if not isinstance(news_articles, list):
        news_articles = []
    
    console.print(f"📊 Company: {company_name} ({ticker})")
    console.print(f"📈 Key Metrics Available: {len(key_metrics) if key_metrics else 0}")
    console.print(f"📰 News Articles: {len(news_articles)}")
    
    # Format stock performance summary
    if key_metrics and isinstance(key_metrics, dict) and any(key_metrics.values()):
        try:
            current_price = key_metrics.get('current_price', 'N/A')
            market_cap = key_metrics.get('market_cap', 'N/A')
            pe_ratio = key_metrics.get('pe_ratio', 'N/A')
            
            if isinstance(market_cap, (int, float)) and market_cap > 0:
                if market_cap >= 1e12:
                    market_cap_str = f"${market_cap/1e12:.2f}T"
                elif market_cap >= 1e9:
                    market_cap_str = f"${market_cap/1e9:.2f}B"
                else:
                    market_cap_str = f"${market_cap/1e6:.2f}M"
            else:
                market_cap_str = str(market_cap)
            
            stock_performance_text = (
                f"Current Price: ${current_price} | "
                f"Market Cap: {market_cap_str} | "
                f"P/E Ratio: {pe_ratio}"
            )
        except Exception as e:
            console.print(f"⚠️ Error formatting metrics: {e}")
            stock_performance_text = "Stock performance data is available but formatting failed."
    else:
        stock_performance_text = "Stock performance data is unavailable."

    if company_name == "Unknown Company":
        error_msg = "Failed to generate a report. Could not determine the company from your query."
        console.print(Panel(error_msg, style="bold red"))
        return {
            "final_answer": error_msg,
            "messages": [AIMessage(content=error_msg)]
        }

    # Format news summary
    if news_articles and isinstance(news_articles, list) and len(news_articles) > 0:
        try:
            news_items = []
            for article in news_articles[:5]:
                if isinstance(article, dict) and article.get('title'):
                    title = str(article.get('title', 'N/A'))
                    if title and title != 'N/A':
                        news_items.append(f"- **{title}**")
            
            if news_items:
                news_summary_text = "\n".join(news_items)
            else:
                news_summary_text = f"- **{company_name}** market developments ongoing\n- Analyst coverage and market sentiment analysis continues"
        except Exception as e:
            console.print(f"⚠️ Error formatting news: {e}")
            news_summary_text = f"- **{company_name}** market developments ongoing"
    else:
        news_summary_text = f"- **{company_name}** market developments ongoing"

    # --- MARKDOWN REPORT CONSTRUCTION ---
    risk_preview = truncate_text(risk_summary, max_length=1000) if risk_summary else "Risk data unavailable."
    
    report_content_md = f"""
# 📊 Financial Analyst Report for {company_name}

## Executive Summary
This report combines key insights from market data, SEC filings, and news sentiment.

## 1. ⚠️ Key Business Risks (from SEC Filings)
{risk_preview}

## 2. 📈 Recent Market Performance
{stock_performance_text}

## 3. 📰 Recent News & Catalysts
{news_summary_text}

---

**Generated:** {datetime.now().strftime("%B %d, %Y at %I:%M %p")}  
**Note:** This report is for informational purposes only.
"""
    
    # Pretty print the report
    console.print("\n" + "="*60)
    console.print(RichMarkdown(report_content_md))
    console.print("="*60)
    
    # Create and save enhanced DOCX
    saved_successfully = False
    docx_path = None
    
    if DOCX_AVAILABLE:
        try:
            reports_dir = './reports'
            os.makedirs(reports_dir, exist_ok=True)
            abs_reports_dir = os.path.abspath(reports_dir)
            console.print(f"📁 Reports directory: {abs_reports_dir}")
            
            if not os.access(reports_dir, os.W_OK):
                console.print(Panel("❌ No write permission to reports directory", style="bold red"))
            else:
                console.print("✅ Write permission confirmed")
            
            console.print("📄 Creating beautiful DOCX document...")
            doc = create_enhanced_docx_report(company_name, risk_summary, stock_performance_text, news_articles, key_metrics)
            
            if doc:
                safe_company_name = "".join(c for c in str(company_name) if c.isalnum() or c in (' ', '-', '_')).strip()
                safe_company_name = safe_company_name.replace(' ', '_')[:50]
                if not safe_company_name:
                    safe_company_name = "Company_Report"
                
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"{safe_company_name}_Report_{timestamp}.docx"
                docx_path = os.path.join(reports_dir, filename)
                
                console.print(f"💾 Saving beautiful document to: {os.path.abspath(docx_path)}")
                
                doc.save(docx_path)
                
                if os.path.exists(docx_path):
                    file_size = os.path.getsize(docx_path)
                    console.print(Panel(f"✅ Beautiful DOCX saved successfully!\nPath: {os.path.abspath(docx_path)}\nSize: {file_size:,} bytes", style="bold green"))
                    saved_successfully = True
                else:
                    console.print(Panel("❌ File was not saved (file does not exist after save)", style="bold red"))
            else:
                console.print(Panel("❌ Document creation failed - doc object is None", style="bold red"))
            
        except Exception as e:
            console.print(Panel(f"❌ DOCX save failed: {str(e)}", style="bold red"))
            import traceback
            console.print(f"Full DOCX error traceback: {traceback.format_exc()}")
    else:
        console.print(Panel("📄 DOCX export skipped (python-docx not available)", style="bold yellow"))

    console.print("---NODE: Beautiful Report Generation Complete---")
    
    try:
        from IPython.display import Markdown, display
        display(Markdown(report_content_md))
    except ImportError:
        pass

    return {
        "report_data": {
            "company_name": company_name,
            "ticker": ticker,
            "risks": risk_summary,
            "performance_metrics": stock_performance_text,
            "news_count": len(news_articles) if isinstance(news_articles, list) else 0,
            "docx_path": docx_path if saved_successfully else None
        },
        "final_answer": report_content_md, 
        "messages": [
            AIMessage(content=f"✨ Here's your comprehensive and beautifully formatted financial analysis report for {company_name} ({ticker}). "
                            f"{'The stunning DOCX file has been generated with enhanced visual styling and can be downloaded.' if saved_successfully else ''}")
        ],
    }